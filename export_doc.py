from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import pymysql

select_type = "select distinct question_type from exam_questions order by question_type asc;"
select_sql = "SELECT id, chapter_no, title, question_type, options, correct_answer, short_answer FROM exam_questions WHERE question_type = %s order by chapter_no;"

conn = pymysql.connect(
    host="localhost",
    user="root",
    password="12345",
    database="user",
    port=3306,
    charset="utf8mb4"
)

cursor = conn.cursor()

# 定义一个函数，统一设置 run 的字体和大小
def set_font(paragraph, font_name="微软雅黑", font_size=14):
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 中文
        run.font.size = Pt(font_size)

# 新建一个 Word 文档
doc = Document()
doc.add_heading("数字化转型初级认证练习题库", level=0)
set_font(doc.paragraphs[-1])

def export_excel(path):
    cursor.execute(select_type)
    ques_type = cursor.fetchall()

    if not ques_type:
        print("未找到题目类型")
        return

    for i, v in enumerate(ques_type):
        cursor.execute(select_sql, v)
        res = cursor.fetchall()

        if v[0] == 'dan_xuan':
            # 单选题标题
            p = doc.add_heading("单选题", level=1)
            set_font(p)

            for j, k in enumerate(res, 1):
                # 题目
                p = doc.add_paragraph(f"{j}. {k[2]}")
                set_font(p)

                # 选项
                options = k[4].strip().replace("\n", "").split("<|>")
                for u in options:
                    p = doc.add_paragraph(f"{u}", style="List Bullet")
                    set_font(p)

                # 答案
                tag = "对" if k[5][-1] == "d" else "错"
                answer = k[5][:-1] + "(" + tag + ")"
                p = doc.add_paragraph(f"正确答案: {answer}", style="Intense Quote")
                set_font(p)

        elif v[0] == 'duo_xuan':
            # 多选题标题
            p = doc.add_heading("多选题", level=1)
            set_font(p)

            for j, k in enumerate(res, 1):
                # 题目
                p = doc.add_paragraph(f"{j}. {k[2]}")
                set_font(p)

                # 选项
                options = k[4].strip().replace("\n", "").split("<|>")
                for u in options:
                    p = doc.add_paragraph(f"{u}", style="List Bullet")
                    set_font(p)

                # 答案
                tag = "对" if k[5][-1] == "d" else "错"
                answer = k[5][:-1] + "(" + tag + ")"
                p = doc.add_paragraph(f"正确答案: {answer}", style="Intense Quote")
                set_font(p)

        elif v[0] == 'pan_duan':
            # 判断题标题
            p = doc.add_heading("判断题", level=1)
            set_font(p)

            for j, k in enumerate(res, 1):
                # 题目
                p = doc.add_paragraph(f"{j}. {k[2]}")
                set_font(p)

                # 答案
                if k[5][-1] == "x":
                    answer = "对" if k[5][:-1] == "错" else "错"
                else:
                    answer = k[5][:-1]

                p = doc.add_paragraph(f"正确答案: {answer}", style="Intense Quote")
                set_font(p)

    # 导出到一个文件
    doc.save("试题导出.docx")

# def export_excel(path):
#     cursor.execute(select_type)
#     ques_type = cursor.fetchall()
#
#     if ques_type is None:
#         print("未找到题目类型")
#         return
#
#     # 根据题目类型导出到excel
#     for i, v in enumerate(ques_type):
#         cursor.execute(select_sql, v)
#         res = cursor.fetchall()
#
#         if v[0] == 'dan_xuan':
#             for j, k in enumerate(res, 1):
#                 # chapter_no = k[1]
#                 # if chapter_no k[1] == :
#                 #     doc1.add_heading(f"{k[1]}", level=3)
#                 # 题目
#                 p = doc1.add_paragraph(f"{j}. {k[2]}")
#                 p.style.font.size = Pt(16)
#
#                 # 选项
#                 options = k[4].strip().replace("\n", "").split("<|>")
#                 for u in options:
#                     doc1.add_paragraph(f"{u}", style="List Bullet")
#
#                 # 正常答案
#                 tag = "对" if k[5][-1] == "d" else "错"
#                 answer = k[5][:-1] + "(" + tag + ")"
#
#                 # 答案
#                 doc1.add_paragraph(f"正确答案: {answer}", style="Intense Quote")
#             doc1.save("单选题v1.0.docx")
#         elif v[0] == 'duo_xuan':
#             for j, k in enumerate(res, 1):
#                 # chapter_no = k[1]
#                 # if chapter_no k[1] == :
#                 #     doc1.add_heading(f"{k[1]}", level=3)
#                 # 题目
#                 p = doc2.add_paragraph(f"{j}. {k[2]}")
#                 p.style.font.size = Pt(16)
#
#                 # 选项
#                 options = k[4].strip().replace("\n", "").split("<|>")
#                 for u in options:
#                     doc2.add_paragraph(f"{u}", style="List Bullet")
#
#                 # 正常答案
#                 tag = "对" if k[5][-1] == "d" else "错"
#                 answer = k[5][:-1] + "(" + tag + ")"
#
#                 # 答案
#                 doc2.add_paragraph(f"正确答案: {answer}", style="Intense Quote")
#             doc2.save("多选题v1.0.docx")
#         elif v[0] == 'pan_duan':
#             for j, k in enumerate(res, 1):
#                 # chapter_no = k[1]
#                 # if chapter_no k[1] == :
#                 #     doc1.add_heading(f"{k[1]}", level=3)
#                 # 题目
#                 p = doc3.add_paragraph(f"{j}. {k[2]}")
#                 p.style.font.size = Pt(16)
#                 # 选项
#                 # options = ' '.join(k[3].strip().replace("\n", "").split("<|>"))
#                 # temp.append(options)
#                 # 正常答案
#                 if k[5][-1] == "x":
#                     answer = "对" if k[5][:-1] == "错" else "错"
#                 else:
#                     answer = k[5][:-1]
#                 # 答案
#                 doc3.add_paragraph(f"正确答案: {answer}", style="Intense Quote")
#             doc3.save("判断题v1.0.docx")


# def format(ws):
#     # 设置整个 sheet 的字体（遍历所有单元格）
#     font = Font(name="微软雅黑", size=18)  # 字体：微软雅黑，11号
#     for row in ws.iter_rows():
#         for cell in row:
#             cell.font = font
#
#     # 设置列宽自动化
#     for col in ws.columns:
#         max_length = 0
#         col_letter = get_column_letter(col[0].column)  # 获取列字母
#         for cell in col:
#             if cell.value:  # 计算每个单元格内容长度
#                 max_length = max(max_length, len(str(cell.value)))
#         ws.column_dimensions[col_letter].width = max_length + 20  # 留点空隙


if __name__ == '__main__':
    export_excel("数字化转型练习题库(不全).xlsx")
